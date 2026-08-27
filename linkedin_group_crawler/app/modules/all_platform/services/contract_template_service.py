"""Thư viện Mẫu hợp đồng — upload file .docx/.pdf/.txt, trích xuất text để AI Contract
Copilot tham chiếu văn phong/cấu trúc khi soạn thảo (contract_ai_service.py).

Chỉ lưu TEXT đã trích xuất trong DB, không lưu file gốc lên Storage — xem lý do trong
migration 074_contract_templates.sql."""

from __future__ import annotations

import io
from datetime import datetime, timezone

from supabase import Client

from app.core.supabase_client import get_supabase_client

TABLE = "contract_templates"

_MAX_TEXT_CHARS = 20000  # đủ dài để AI tham chiếu, tránh prompt quá to/tốn token


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _row_to_template(row: dict, include_text: bool = False) -> dict:
    out = {
        "id": row["id"],
        "name": row["name"],
        "description": row.get("description") or "",
        "fileName": row["file_name"],
        "fileType": row["file_type"],
        "textLength": len(row.get("extracted_text") or ""),
        "createdById": row.get("created_by"),
        "createdAt": row.get("created_at"),
        "updatedAt": row.get("updated_at"),
    }
    if include_text:
        out["extractedText"] = row.get("extracted_text") or ""
    return out


def extract_text_from_file(file_name: str, content: bytes) -> tuple[str, str]:
    """Trích xuất text từ file mẫu hợp đồng. Trả về (file_type, text).
    Raise ValueError nếu định dạng không hỗ trợ hoặc file lỗi/không đọc được."""
    ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""

    if ext == "docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise ValueError("Thiếu thư viện python-docx trên server.") from exc
        try:
            document = Document(io.BytesIO(content))
        except Exception as exc:
            raise ValueError("Không đọc được file .docx — file có thể bị hỏng.") from exc
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
        return "docx", text[:_MAX_TEXT_CHARS]

    if ext == "pdf":
        try:
            import pdfplumber
        except ImportError as exc:
            raise ValueError("Thiếu thư viện pdfplumber trên server.") from exc
        try:
            parts = []
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        parts.append(page_text)
            text = "\n".join(parts)
        except Exception as exc:
            raise ValueError("Không đọc được file .pdf — file có thể bị hỏng hoặc là ảnh scan (không có text).") from exc
        return "pdf", text[:_MAX_TEXT_CHARS]

    if ext == "txt":
        try:
            text = content.decode("utf-8", errors="replace")
        except Exception as exc:
            raise ValueError("Không đọc được file .txt.") from exc
        return "txt", text[:_MAX_TEXT_CHARS]

    raise ValueError(f"Định dạng .{ext or '?'} chưa được hỗ trợ — chỉ nhận .docx, .pdf, .txt.")


def list_contract_templates() -> list[dict]:
    supabase: Client = get_supabase_client()
    result = supabase.table(TABLE).select("*").order("created_at", desc=True).execute()
    return [_row_to_template(row) for row in (result.data or [])]


def get_contract_template(template_id: str, include_text: bool = True) -> dict:
    supabase: Client = get_supabase_client()
    row = supabase.table(TABLE).select("*").eq("id", template_id).single().execute().data
    return _row_to_template(row, include_text=include_text)


def create_contract_template(name: str, description: str | None, file_name: str, content: bytes, created_by: str | None) -> dict:
    if not content:
        raise ValueError("File rỗng, không có nội dung.")
    file_type, text = extract_text_from_file(file_name, content)
    if not text.strip():
        raise ValueError("Không trích xuất được nội dung text nào từ file — file có thể là ảnh scan hoặc rỗng.")
    supabase: Client = get_supabase_client()
    insert_data = {
        "name": name,
        "description": description,
        "file_name": file_name,
        "file_type": file_type,
        "extracted_text": text,
        "created_by": created_by,
    }
    row = supabase.table(TABLE).insert(insert_data).execute().data[0]
    return _row_to_template(row)


def delete_contract_template(template_id: str) -> None:
    supabase: Client = get_supabase_client()
    supabase.table(TABLE).delete().eq("id", template_id).execute()
